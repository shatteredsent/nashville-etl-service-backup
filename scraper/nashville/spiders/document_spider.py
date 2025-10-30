"""
Document Spider for CSV, Excel, and Word Documents
Handles: .csv, .xlsx, .xls, .docx files
Follows: Single Responsibility, Clean Architecture, SOLID Principles
"""
import os
import re
import hashlib
from typing import Dict, Any, List, Optional, Union
from pathlib import Path

import scrapy
import pandas as pd
from docx import Document

from scraper.nashville.items import BusinessItem


class DocumentSpider(scrapy.Spider):
    """Spider for processing structured documents (CSV, Excel, Word)."""

    name = 'document'

    # Column name mappings for flexible parsing
    COLUMN_MAPPINGS = {
        'name': ['name', 'event_name', 'title', 'event', 'business_name'],
        'venue_name': ['venue', 'venue_name', 'location', 'place'],
        'venue_address': ['address', 'venue_address', 'street_address', 'location_address'],
        'event_date': ['date', 'event_date', 'datetime', 'when', 'event_time'],
        'description': ['description', 'details', 'info', 'about'],
        'url': ['url', 'website', 'link', 'web'],
        'category': ['category', 'type', 'genre', 'event_type'],
    }

    def __init__(self, file_path: Optional[str] = None, *args, **kwargs):
        """
        Initialize spider with file path.

        Args:
            file_path: Absolute path to the document to process

        Raises:
            ValueError: If file_path is not provided or invalid
        """
        super().__init__(*args, **kwargs)
        self._validate_initialization(file_path)
        self.file_path = file_path
        self.file_extension = self._get_file_extension(file_path)

    def _validate_initialization(self, file_path: Optional[str]) -> None:
        """Validate spider initialization parameters."""
        if not file_path:
            raise ValueError("file_path argument is required")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document file not found: {file_path}")

    def _get_file_extension(self, file_path: str) -> str:
        """Extract and validate file extension."""
        extension = Path(file_path).suffix.lower()
        supported = {'.csv', '.xlsx', '.xls', '.docx'}
        if extension not in supported:
            raise ValueError(
                f"Unsupported file type: {extension}. Supported: {supported}")
        return extension

    def start_requests(self):
        """Initiate spider request."""
        yield scrapy.Request(
            url=f'file://{os.path.abspath(self.file_path)}',
            callback=self.parse,
            dont_filter=True
        )

    def parse(self, response):
        """
        Main parsing dispatcher.

        Args:
            response: Scrapy response object

        Yields:
            BusinessItem objects
        """
        try:
            items = self._extract_items_by_type()
            valid_items = self._validate_items(items)

            self.logger.info(
                f"Extracted {len(items)} items, {len(valid_items)} valid")

            for item_data in valid_items:
                yield self._create_business_item(item_data)

        except Exception as e:
            self.logger.error(f"Parse failed for {self.file_path}: {e}")
            raise

    def _extract_items_by_type(self) -> List[Dict[str, Any]]:
        """
        Route extraction based on file type.

        Returns:
            List of extracted item dictionaries
        """
        extractors = {
            '.csv': self._extract_from_csv,
            '.xlsx': self._extract_from_excel,
            '.xls': self._extract_from_excel,
            '.docx': self._extract_from_word,
        }

        extractor = extractors.get(self.file_extension)
        if not extractor:
            raise ValueError(f"No extractor for {self.file_extension}")

        return extractor()

    def _extract_from_csv(self) -> List[Dict[str, Any]]:
        """
        Extract data from CSV file.

        Returns:
            List of item dictionaries
        """
        try:
            df = pd.read_csv(self.file_path, encoding='utf-8')
        except UnicodeDecodeError:
            try:
                df = pd.read_csv(self.file_path, encoding='latin-1')
            except Exception as e:
                self.logger.error(f"CSV encoding error: {e}")
                return []

        return self._dataframe_to_items(df)

    def _extract_from_excel(self) -> List[Dict[str, Any]]:
        """
        Extract data from Excel file.

        Returns:
            List of item dictionaries
        """
        try:
            # Try reading first sheet
            df = pd.read_excel(self.file_path, sheet_name=0)
            items = self._dataframe_to_items(df)

            # If no valid items, try all sheets
            if not items:
                xls = pd.ExcelFile(self.file_path)
                all_items = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(self.file_path, sheet_name=sheet_name)
                    all_items.extend(self._dataframe_to_items(df))
                return all_items

            return items

        except Exception as e:
            self.logger.error(f"Excel extraction error: {e}")
            return []

    def _extract_from_word(self) -> List[Dict[str, Any]]:
        """
        Extract data from Word document.

        Returns:
            List of item dictionaries
        """
        try:
            doc = Document(self.file_path)

            # Try extracting from tables first
            items = self._extract_from_word_tables(doc)

            # If no tables or no valid items, parse text
            if not items:
                items = self._extract_from_word_text(doc)

            return items

        except Exception as e:
            self.logger.error(f"Word extraction error: {e}")
            return []

    def _extract_from_word_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract data from Word document tables.

        Args:
            doc: python-docx Document object

        Returns:
            List of item dictionaries
        """
        items = []

        for table in doc.tables:
            try:
                # Convert table to list of lists
                data = [[cell.text.strip() for cell in row.cells]
                        for row in table.rows]

                if len(data) < 2:  # Need header + at least one row
                    continue

                # Create DataFrame from table
                df = pd.DataFrame(data[1:], columns=data[0])
                items.extend(self._dataframe_to_items(df))

            except Exception as e:
                self.logger.debug(f"Table extraction error: {e}")
                continue

        return items

    def _extract_from_word_text(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract data from Word document paragraphs.

        Args:
            doc: python-docx Document object

        Returns:
            List of item dictionaries
        """
        items = []
        current_item = {}

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text or len(text) < 3:
                # End of current item
                if current_item.get('name'):
                    items.append(current_item.copy())
                    current_item = {}
                continue

            # Parse structured data (Key: Value format)
            if ':' in text and not text.startswith('http'):
                key, value = self._parse_key_value(text)
                if key and value:
                    current_item[key] = value
            else:
                # Classify line content
                self._classify_text_line(text, current_item)

        # Add final item if exists
        if current_item.get('name'):
            items.append(current_item)

        return items

    def _dataframe_to_items(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """
        Convert DataFrame to list of item dictionaries.

        Args:
            df: pandas DataFrame

        Returns:
            List of item dictionaries
        """
        if df.empty:
            return []

        # Normalize column names
        df = self._normalize_dataframe_columns(df)

        # Convert to dictionary records
        items = df.to_dict('records')

        # Clean items
        return [self._clean_item(item) for item in items]

    def _normalize_dataframe_columns(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Normalize DataFrame column names to standard fields.

        Args:
            df: pandas DataFrame

        Returns:
            DataFrame with normalized column names
        """
        # Create lowercase mapping
        columns_lower = {col: col.lower().strip() for col in df.columns}
        df.rename(columns=columns_lower, inplace=True)

        # Map to standard fields
        rename_map = {}
        for standard_name, alternatives in self.COLUMN_MAPPINGS.items():
            for col in df.columns:
                if col in alternatives:
                    rename_map[col] = standard_name
                    break

        df.rename(columns=rename_map, inplace=True)

        return df

    def _parse_key_value(self, text: str) -> tuple[Optional[str], Optional[str]]:
        """
        Parse key-value pair from text.

        Args:
            text: Text containing key: value format

        Returns:
            Tuple of (key, value)
        """
        parts = text.split(':', 1)
        if len(parts) != 2:
            return None, None

        key = parts[0].strip().lower()
        value = parts[1].strip()

        # Map key to standard field
        for standard_name, alternatives in self.COLUMN_MAPPINGS.items():
            if key in alternatives:
                return standard_name, value

        return key, value

    def _classify_text_line(self, text: str, item: Dict[str, Any]) -> None:
        """
        Classify and add text line to item.

        Args:
            text: Text line to classify
            item: Item dictionary to update (modified in place)
        """
        # Check for URL
        if self._is_url(text):
            item['url'] = text
        # Check for date
        elif self._is_date(text):
            item['event_date'] = text
        # Check for address
        elif self._is_address(text):
            item['venue_address'] = text
        # Check if it looks like a name/title
        elif self._looks_like_name(text):
            if not item.get('name'):
                item['name'] = text
                item['venue_name'] = text
            else:
                # Add to description
                item.setdefault('description', []).append(text)
        else:
            # Add to description
            item.setdefault('description', []).append(text)

    def _is_url(self, text: str) -> bool:
        """Check if text is a URL."""
        return bool(re.match(r'https?://', text))

    def _is_date(self, text: str) -> bool:
        """Check if text looks like a date."""
        date_patterns = [
            r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
            r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b',
            r'\b\d{4}-\d{2}-\d{2}\b',
        ]
        return any(re.search(pattern, text.lower()) for pattern in date_patterns)

    def _is_address(self, text: str) -> bool:
        """Check if text looks like an address."""
        address_keywords = ['street', 'st', 'avenue', 'ave', 'road', 'rd',
                            'boulevard', 'blvd', 'drive', 'dr', 'nashville', 'tn']
        return any(keyword in text.lower() for keyword in address_keywords)

    def _looks_like_name(self, text: str) -> bool:
        """Check if text looks like an event/business name."""
        if not (5 <= len(text) <= 150):
            return False
        if not text[0].isupper():
            return False
        return True

    def _clean_item(self, item: Dict[str, Any]) -> Dict[str, Any]:
        """
        Clean and standardize item data.

        Args:
            item: Raw item dictionary

        Returns:
            Cleaned item dictionary
        """
        cleaned = {}

        # Handle description list
        if isinstance(item.get('description'), list):
            item['description'] = ' '.join(item['description'])[:500]

        # Clean each field
        for key, value in item.items():
            if pd.isna(value) or value == '' or value is None:
                continue

            # Convert to string and strip
            cleaned[key] = str(value).strip()

        return cleaned

    def _validate_items(self, items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Validate and filter items.

        Args:
            items: List of item dictionaries

        Returns:
            List of valid items
        """
        valid_items = []

        for item in items:
            if self._is_valid_item(item):
                valid_items.append(item)
            else:
                self.logger.debug(
                    f"Skipping invalid item: {item.get('name', 'Unknown')}")

        return valid_items

    def _is_valid_item(self, item: Dict[str, Any]) -> bool:
        """
        Check if item meets minimum requirements.

        Args:
            item: Item dictionary

        Returns:
            True if valid, False otherwise
        """
        # Must have a name
        name = item.get('name', '')
        if not name or len(name) < 3:
            return False

        # Name shouldn't be just numbers or symbols
        if not any(c.isalpha() for c in name):
            return False

        return True

    def _create_business_item(self, data: Dict[str, Any]) -> BusinessItem:
        """
        Create BusinessItem from data dictionary.

        Args:
            data: Item data dictionary

        Returns:
            BusinessItem object
        """
        item = BusinessItem()

        # Required fields
        item['source'] = self._get_source_name()
        item['name'] = data.get('name', '').strip()

        # Optional fields
        item['venue_name'] = data.get(
            'venue_name', data.get('name', '')).strip()
        item['venue_address'] = data.get('venue_address', '').strip()
        item['venue_city'] = data.get('venue_city', 'Nashville')
        item['description'] = data.get('description', '').strip()
        item['event_date'] = data.get('event_date')
        item['category'] = data.get('category', 'document_extracted')

        # Generate or use URL
        item['url'] = self._get_or_generate_url(data)

        return item

    def _get_source_name(self) -> str:
        """Get display name for data source."""
        filename = os.path.basename(self.file_path)
        return f"document_upload_{self.file_extension[1:]}"

    def _get_or_generate_url(self, data: Dict[str, Any]) -> str:
        """
        Get existing URL or generate unique identifier.

        Args:
            data: Item data dictionary

        Returns:
            URL string
        """
        # Use existing URL if present
        url = data.get('url', '').strip()
        if url and len(url) > 5 and url.startswith('http'):
            return url

        # Generate unique URL
        content = f"{data.get('name', '')}|{data.get('venue_address', '')}|{os.path.basename(self.file_path)}"
        hash_value = hashlib.md5(content.encode()).hexdigest()[:12]

        return f"document://{self.file_extension[1:]}-event/{hash_value}"
