"""
Word document extractor.
Handles .docx files with text and tables.
"""
import re
from typing import List, Dict, Any
from docx import Document
from .base_extractor import BaseExtractor


class WordExtractor(BaseExtractor):
    """Extract structured data from Word documents."""

    SUPPORTED_EXTENSIONS = {'.docx'}

    # Patterns for identifying data types
    PHONE_PATTERN = re.compile(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
    URL_PATTERN = re.compile(r'https?://[^\s]+')
    DATE_PATTERNS = [
        re.compile(r'\d{1,2}/\d{1,2}/\d{2,4}'),
        re.compile(r'\d{4}-\d{2}-\d{2}'),
        re.compile(
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]* \d{1,2}', re.IGNORECASE),
    ]

    ADDRESS_INDICATORS = frozenset([
        'street', 'st', 'avenue', 'ave', 'road', 'rd',
        'boulevard', 'blvd', 'drive', 'dr', 'nashville', 'tn'
    ])

    def extract(self) -> List[Dict[str, Any]]:
        """Extract structured data from Word document."""
        try:
            doc = Document(str(self.file_path))

            # Extract from paragraphs
            paragraph_items = self._extract_from_paragraphs(doc.paragraphs)
            self._log(
                f"Extracted {len(paragraph_items)} items from paragraphs")

            # Extract from tables
            table_items = self._extract_from_tables(doc.tables)
            self._log(f"Extracted {len(table_items)} items from tables")

            all_items = paragraph_items + table_items
            self._log(f"Total {len(all_items)} items from Word document")

            return [self._create_item(**item) for item in all_items if self._is_valid(item)]

        except Exception as e:
            self._log(f"Word document extraction failed: {e}", 'error')
            raise

    def _extract_from_paragraphs(self, paragraphs) -> List[Dict[str, Any]]:
        """Extract data from document paragraphs."""
        lines = [p.text.strip() for p in paragraphs if p.text.strip()]

        if not lines:
            return []

        return self._parse_lines(lines)

    def _extract_from_tables(self, tables) -> List[Dict[str, Any]]:
        """Extract data from document tables."""
        all_items = []

        for table in tables:
            items = self._process_table(table)
            all_items.extend(items)

        return all_items

    def _process_table(self, table) -> List[Dict[str, Any]]:
        """Process a single table."""
        if not table.rows:
            return []

        # Try to extract header
        header_row = table.rows[0]
        header = [cell.text.strip().lower() for cell in header_row.cells]

        # Map columns
        column_mapping = self._map_table_columns(header)

        if not column_mapping:
            # No recognizable structure, treat as plain text
            return []

        # Extract data rows
        items = []
        for row in table.rows[1:]:
            item = self._parse_table_row(row, column_mapping)
            if item and item.get('name'):
                items.append(item)

        return items

    def _map_table_columns(self, header: List[str]) -> Dict[str, int]:
        """Map table columns to field names."""
        name_cols = {'name', 'event', 'venue', 'title'}
        address_cols = {'address', 'location', 'venue_address'}
        date_cols = {'date', 'event_date', 'when'}
        desc_cols = {'description', 'details', 'info'}
        url_cols = {'url', 'website', 'link'}

        mapping = {}
        for idx, col in enumerate(header):
            if col in name_cols:
                mapping['name'] = idx
            elif col in address_cols:
                mapping['venue_address'] = idx
            elif col in date_cols:
                mapping['event_date'] = idx
            elif col in desc_cols:
                mapping['description'] = idx
            elif col in url_cols:
                mapping['url'] = idx

        return mapping

    def _parse_table_row(self, row, column_mapping: Dict[str, int]) -> Dict[str, Any]:
        """Parse a table row into an item."""
        item = {}
        cells = [cell.text.strip() for cell in row.cells]

        for field, col_idx in column_mapping.items():
            if col_idx < len(cells) and cells[col_idx]:
                item[field] = cells[col_idx]

        if item.get('name'):
            item['venue_name'] = item['name']

        return item

    def _parse_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Parse lines into structured items (similar to PDF)."""
        items = []
        current_item = {}

        for line in lines:
            if ':' in line and not line.startswith('http'):
                current_item = self._handle_structured_line(
                    line, current_item, items)
            else:
                current_item = self._handle_unstructured_line(
                    line, current_item, items)

        if current_item and current_item.get('name'):
            items.append(current_item)

        return items

    def _handle_structured_line(
        self,
        line: str,
        current_item: Dict,
        items: List[Dict]
    ) -> Dict:
        """Process structured 'Label: Value' line."""
        parts = line.split(':', 1)
        label = parts[0].strip().lower()
        value = parts[1].strip() if len(parts) > 1 else ''

        if label in ['venue', 'location', 'place', 'event', 'name']:
            if current_item and current_item.get('name'):
                items.append(current_item)
            return {'name': value, 'venue_name': value}

        field_map = {
            'address': 'venue_address',
            'date': 'event_date',
            'phone': 'phone',
            'url': 'url',
            'website': 'url',
        }

        if label in field_map:
            current_item[field_map[label]] = value
        else:
            self._append_to_description(current_item, line)

        return current_item

    def _handle_unstructured_line(
        self,
        line: str,
        current_item: Dict,
        items: List[Dict]
    ) -> Dict:
        """Process unstructured line."""
        line_type = self._identify_line_type(line)

        if line_type == 'name':
            if current_item and current_item.get('name'):
                items.append(current_item)
            return {'name': line, 'venue_name': line}

        if line_type == 'address':
            current_item['venue_address'] = line
        elif line_type == 'date':
            current_item['event_date'] = line
        elif line_type == 'url':
            current_item['url'] = line
        else:
            self._append_to_description(current_item, line)

        return current_item

    def _identify_line_type(self, line: str) -> str:
        """Identify line type."""
        if self.URL_PATTERN.search(line):
            return 'url'
        if any(pattern.search(line) for pattern in self.DATE_PATTERNS):
            return 'date'
        if any(ind in line.lower() for ind in self.ADDRESS_INDICATORS):
            return 'address'
        if self._is_likely_name(line):
            return 'name'
        return 'text'

    def _is_likely_name(self, text: str) -> bool:
        """Check if text looks like a business/event name."""
        if not (5 <= len(text) <= 100):
            return False
        return text[0].isupper() or any(
            word in text.lower()
            for word in ['&', 'and', "'s", 'the ', 'venue']
        )

    def _append_to_description(self, item: Dict, text: str) -> None:
        """Append text to description."""
        if 'description' not in item:
            item['description'] = text
        else:
            item['description'] += ' ' + text

    def _is_valid(self, item: Dict) -> bool:
        """Validate item."""
        return bool(item.get('name') and len(item['name']) >= 3)
